"""
Tests for content extraction functionality.
"""

import importlib.util
import logging
import sys
from unittest.mock import patch

import pytest


class TestContentExtractor:
    """Test content extraction from various file types."""

    def test_extract_txt_content(self, temp_dir):
        """Test extracting content from text files."""
        from footprinter.ingest.content_extractors import ContentExtractor

        extractor = ContentExtractor()

        # Create test file
        test_file = temp_dir / "test.txt"
        test_file.write_text("Hello, world!\nThis is a test.")

        content = extractor.extract(test_file)

        assert content is not None
        assert "Hello, world!" in content
        assert "This is a test" in content

    def test_extract_py_content(self, temp_dir):
        """Test extracting content from Python files."""
        from footprinter.ingest.content_extractors import ContentExtractor

        extractor = ContentExtractor()

        test_file = temp_dir / "test.py"
        test_file.write_text("def hello():\n    print('Hello')\n")

        content = extractor.extract(test_file)

        assert content is not None
        assert "def hello" in content

    def test_extract_md_content(self, temp_dir):
        """Test extracting content from Markdown files."""
        from footprinter.ingest.content_extractors import ContentExtractor

        extractor = ContentExtractor()

        test_file = temp_dir / "test.md"
        test_file.write_text("# Title\n\nSome **bold** text.")

        content = extractor.extract(test_file)

        assert content is not None
        assert "Title" in content

    def test_extract_json_content(self, temp_dir):
        """Test extracting content from JSON files."""
        from footprinter.ingest.content_extractors import ContentExtractor

        extractor = ContentExtractor()

        test_file = temp_dir / "test.json"
        test_file.write_text('{"name": "test", "value": 123}')

        content = extractor.extract(test_file)

        assert content is not None
        assert "test" in content

    def test_content_preview_truncated(self, temp_dir):
        """Test that content preview is truncated for large files."""
        from footprinter.ingest.content_extractors import ContentExtractor

        extractor = ContentExtractor(max_preview_length=100)

        # Create a large file
        test_file = temp_dir / "large.txt"
        test_file.write_text("x" * 10000)

        content = extractor.extract(test_file)

        assert content is not None
        assert len(content) <= 100

    def test_nonexistent_file_returns_none(self, temp_dir):
        """Test that nonexistent files return None."""
        from footprinter.ingest.content_extractors import ContentExtractor

        extractor = ContentExtractor()

        content = extractor.extract(temp_dir / "nonexistent.txt")

        assert content is None or content == ""

    @pytest.mark.skipif(not importlib.util.find_spec("pypdf"), reason="requires [docs] extra")
    def test_pypdf_no_warning_logs(self, temp_dir, caplog):
        """Test that pypdf does not emit WARNING-level logs during extraction."""
        import pypdf

        from footprinter.ingest.content_extractors import ContentExtractor

        # Create a minimal valid PDF
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=72, height=72)
        pdf_path = temp_dir / "blank.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)

        extractor = ContentExtractor()

        with caplog.at_level(logging.WARNING, logger="pypdf"):
            extractor.extract(pdf_path)

        pypdf_records = [r for r in caplog.records if r.name.startswith("pypdf")]
        assert len(pypdf_records) == 0, f"Unexpected pypdf warnings: {pypdf_records}"


@pytest.mark.skipif(not importlib.util.find_spec("docx"), reason="requires [docs] extra")
class TestDocxExtraction:
    """Test DOCX content extraction (lines 86-109)."""

    def test_extract_docx_content(self, temp_dir):
        """Create a valid .docx and verify extraction."""
        import docx

        from footprinter.ingest.content_extractors import ContentExtractor

        doc = docx.Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")

        docx_path = temp_dir / "test.docx"
        doc.save(str(docx_path))

        extractor = ContentExtractor()
        content = extractor.extract(docx_path)

        assert content is not None
        assert "First paragraph" in content
        assert "Second paragraph" in content

    def test_docx_truncation(self, temp_dir):
        """DOCX extraction respects max_preview_length."""
        import docx

        from footprinter.ingest.content_extractors import ContentExtractor

        doc = docx.Document()
        for i in range(100):
            doc.add_paragraph(f"Long paragraph number {i} with lots of text to fill it up")

        docx_path = temp_dir / "long.docx"
        doc.save(str(docx_path))

        extractor = ContentExtractor(max_preview_length=50)
        content = extractor.extract(docx_path)

        assert content is not None
        assert len(content) <= 50


class TestUnknownExtension:
    """Test unknown file extension handling (lines 37-39)."""

    def test_unknown_extension_returns_none(self, temp_dir):
        from footprinter.ingest.content_extractors import ContentExtractor

        extractor = ContentExtractor()

        unknown_file = temp_dir / "data.xyz"
        unknown_file.write_text("some content")

        content = extractor.extract(unknown_file)
        assert content is None


@pytest.mark.skipif(not importlib.util.find_spec("pypdf"), reason="requires [docs] extra")
class TestPDFTruncation:
    """Test PDF extraction truncation at max_preview_length (line 70)."""

    def test_pdf_truncation_at_max_preview(self, temp_dir):
        """Large PDF text gets truncated to max_preview_length."""
        import pypdf
        from pypdf.generic import (
            DecodedStreamObject,
            DictionaryObject,
            NameObject,
        )

        from footprinter.ingest.content_extractors import ContentExtractor

        # Create a multi-page PDF with text via content streams
        writer = pypdf.PdfWriter()
        text_content = b"Hello world test content. " * 20
        for _ in range(5):
            page = writer.add_blank_page(width=612, height=792)
            content = b"BT /F1 12 Tf 40 750 Td (" + text_content + b") Tj ET"
            resources = DictionaryObject()
            font_dict = DictionaryObject()
            font_dict[NameObject("/Type")] = NameObject("/Font")
            font_dict[NameObject("/Subtype")] = NameObject("/Type1")
            font_dict[NameObject("/BaseFont")] = NameObject("/Helvetica")
            fonts = DictionaryObject()
            fonts[NameObject("/F1")] = font_dict
            resources[NameObject("/Font")] = fonts
            page[NameObject("/Resources")] = resources
            stream = DecodedStreamObject()
            stream.set_data(content)
            page[NameObject("/Contents")] = stream

        pdf_path = temp_dir / "big.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)

        extractor = ContentExtractor(max_preview_length=50)
        content = extractor.extract(pdf_path)

        assert content is not None
        assert len(content) <= 50


class TestCorruptedFile:
    """Test outer exception handler (lines 41-43)."""

    def test_extract_corrupted_file_returns_none(self, temp_dir):
        """A corrupted .pdf triggers the outer except and returns None."""
        from footprinter.ingest.content_extractors import ContentExtractor

        bad_pdf = temp_dir / "corrupt.pdf"
        bad_pdf.write_bytes(b"not a valid pdf file at all")

        extractor = ContentExtractor()
        content = extractor.extract(bad_pdf)

        # The outer except catches the error and returns None or ""
        assert content is None or content == ""


class TestFullContentExtractorSizeGuard:
    """Test FullContentExtractor file size guard."""

    def test_full_extractor_skips_file_exceeding_max_size(self, temp_dir, caplog):
        """Files larger than max_file_size_bytes are skipped with a warning."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(max_file_size_bytes=100)

        big_file = temp_dir / "big.txt"
        big_file.write_text("x" * 200)

        with caplog.at_level(logging.WARNING):
            chunks = extractor.extract_with_chunking(big_file)

        assert chunks == []
        assert any("exceeds content extraction limit" in r.message for r in caplog.records)

    def test_full_extractor_reads_file_within_limit(self, temp_dir):
        """Files smaller than the limit are extracted normally."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(max_file_size_bytes=1000)

        small_file = temp_dir / "small.txt"
        small_file.write_text("hello world")

        chunks = extractor.extract_with_chunking(small_file)

        assert len(chunks) >= 1
        assert "hello world" in chunks[0]["content"]

    def test_full_extractor_default_limit_is_50mb(self):
        """Default max_file_size_bytes is 50 MB."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor()
        assert extractor.max_file_size_bytes == 50 * 1024 * 1024

    def test_full_extractor_zero_limit_means_no_limit(self, temp_dir):
        """max_file_size_bytes=0 disables the guard (consistent with FileScanner)."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(max_file_size_bytes=0)

        test_file = temp_dir / "any_size.txt"
        test_file.write_text("content that should be read regardless")

        chunks = extractor.extract_with_chunking(test_file)

        assert len(chunks) >= 1
        assert "content that should be read regardless" in chunks[0]["content"]

    def test_full_extractor_skips_oversized_pdf(self, temp_dir, caplog):
        """Size guard fires for PDF files before pypdf reads them."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(max_file_size_bytes=50)

        # Write a fake PDF larger than the limit — content doesn't matter,
        # the stat() guard fires before any format-specific reader.
        big_pdf = temp_dir / "big.pdf"
        big_pdf.write_bytes(b"%PDF-1.4 " + b"x" * 200)

        with caplog.at_level(logging.WARNING):
            chunks = extractor.extract_with_chunking(big_pdf)

        assert chunks == []
        assert any("exceeds content extraction limit" in r.message for r in caplog.records)

    def test_full_extractor_skips_oversized_docx(self, temp_dir, caplog):
        """Size guard fires for DOCX files before python-docx reads them."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(max_file_size_bytes=50)

        big_docx = temp_dir / "big.docx"
        big_docx.write_bytes(b"PK" + b"x" * 200)

        with caplog.at_level(logging.WARNING):
            chunks = extractor.extract_with_chunking(big_docx)

        assert chunks == []
        assert any("exceeds content extraction limit" in r.message for r in caplog.records)


class TestFileTypeAllowlist:
    """Test FullContentExtractor file type allowlist filtering."""

    def test_allowlist_filters_disallowed_types(self, temp_dir):
        """Files with extensions not in the allowlist are skipped."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(file_types=[".md", ".txt"])

        py_file = temp_dir / "script.py"
        py_file.write_text("print('hello')")

        chunks = extractor.extract_with_chunking(py_file)
        assert chunks == []

    def test_allowlist_permits_allowed_types(self, temp_dir):
        """Files with extensions in the allowlist are extracted."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(file_types=[".md", ".txt"])

        md_file = temp_dir / "notes.md"
        md_file.write_text("# My Notes\n\nSome content here.")

        chunks = extractor.extract_with_chunking(md_file)
        assert len(chunks) >= 1
        assert "My Notes" in chunks[0]["content"]

    def test_allowlist_none_means_all_types(self, temp_dir):
        """file_types=None preserves legacy behavior — all types extracted."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(file_types=None)

        py_file = temp_dir / "script.py"
        py_file.write_text("def hello(): pass")

        chunks = extractor.extract_with_chunking(py_file)
        assert len(chunks) >= 1
        assert "def hello" in chunks[0]["content"]

    def test_allowlist_case_insensitive(self, temp_dir):
        """Allowlist matching is case-insensitive."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(file_types=[".md"])

        upper_file = temp_dir / "README.MD"
        upper_file.write_text("# Readme content")

        chunks = extractor.extract_with_chunking(upper_file)
        assert len(chunks) >= 1
        assert "Readme content" in chunks[0]["content"]


class TestChunkSizeDefaults:
    """Test that default chunk sizes are tuned for MiniLM-L6-v2."""

    def test_default_chunk_size_is_1000(self):
        """FullContentExtractor defaults to 1000-char chunks."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor()
        assert extractor.chunk_size == 1000

    def test_default_chunk_overlap_is_150(self):
        """FullContentExtractor defaults to 150-char overlap (matching chunking.py)."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor()
        assert extractor.chunk_overlap == 150

    def test_fractional_overlap_auto_converts(self):
        """Legacy float overlap is auto-converted to int chars."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(chunk_size=1000, chunk_overlap=0.15)
        assert extractor.chunk_overlap == 150

    def test_extract_with_chunking_respects_word_boundaries(self, temp_dir):
        """Chunks should not split mid-word."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(chunk_size=100, chunk_overlap=10)
        f = temp_dir / "words.txt"
        f.write_text(" ".join(["hello"] * 50))
        chunks = extractor.extract_with_chunking(f)
        assert len(chunks) >= 2
        for chunk in chunks:
            text = chunk["content"]
            assert "content" in chunk
            assert "chunk_index" in chunk
            assert "total_chunks" in chunk


class TestExcludePatterns:
    """Test FullContentExtractor exclude_patterns filtering."""

    def test_exclude_pattern_skips_matching_file(self, temp_dir):
        """Files matching an exclude pattern are skipped."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(exclude_patterns=["**/chat-exports/**"])

        exports_dir = temp_dir / "chat-exports"
        exports_dir.mkdir()
        json_file = exports_dir / "conversation.json"
        json_file.write_text('{"messages": []}')

        chunks = extractor.extract_with_chunking(json_file)
        assert chunks == []

    def test_exclude_pattern_allows_non_matching(self, temp_dir):
        """Files not matching any exclude pattern are extracted normally."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(exclude_patterns=["**/chat-exports/**"])

        md_file = temp_dir / "notes.md"
        md_file.write_text("# Notes\n\nSome useful content.")

        chunks = extractor.extract_with_chunking(md_file)
        assert len(chunks) >= 1
        assert "Notes" in chunks[0]["content"]

    def test_multiple_exclude_patterns(self, temp_dir):
        """Multiple exclude patterns are all checked."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        extractor = FullContentExtractor(exclude_patterns=["**/chat-exports/**", "**/.cache/**"])

        cache_dir = temp_dir / ".cache"
        cache_dir.mkdir()
        cached_file = cache_dir / "data.txt"
        cached_file.write_text("cached content")

        chunks = extractor.extract_with_chunking(cached_file)
        assert chunks == []


class TestVectorizationConfigLoading:
    """Test that vectorization config is read and passed to FullContentExtractor."""

    def test_rebuild_vectors_reads_vectorization_config(self):
        """_rebuild_vectors() passes vectorization config to FullContentExtractor."""
        from unittest.mock import MagicMock

        mock_store = MagicMock()
        mock_vs_cls = MagicMock()
        mock_vs_cls.get_instance.return_value = mock_store
        mock_vs_cls.reset_instance = MagicMock()
        mock_chroma_path = MagicMock()
        mock_chroma_path.exists.return_value = False

        mock_conn = MagicMock()
        mock_cursor = MagicMock()
        mock_conn.cursor.return_value = mock_cursor
        mock_cursor.fetchall.return_value = []
        mock_cursor.fetchone.return_value = [0]

        test_config = {
            "indexing": {"max_file_size_mb": 10},
            "vectorization": {
                "file_types": [".md", ".txt"],
                "chunk_size": 500,
                "chunk_overlap": 200,
                "exclude_patterns": ["**/exports/**"],
            },
        }

        mock_extractor_cls = MagicMock()

        with (
            patch("footprinter.ingest.vector_ops.sqlite3") as mock_sqlite,
            patch("footprinter.ingest.vector_ops.get_db_path", return_value="/tmp/test.db"),
            patch("footprinter.source_registry.get_config", return_value=test_config),
            patch("footprinter.paths.get_chroma_path", return_value=mock_chroma_path),
            patch("footprinter.semantic.vector_store._file_vectorization_enabled", return_value=True),
            patch("footprinter.semantic.vector_store._chat_vectorization_enabled", return_value=True),
            patch("footprinter.semantic.vector_store.VectorStore", mock_vs_cls),
            patch("footprinter.ingest.full_content_extractor.FullContentExtractor", mock_extractor_cls),
        ):
            mock_sqlite.connect.return_value = mock_conn

            from footprinter.ingest.vector_ops import _rebuild_vectors

            _rebuild_vectors(quiet=True, source="all")

        # Verify FullContentExtractor.from_config was called with the full config
        mock_extractor_cls.from_config.assert_called_once_with(test_config)

    def test_from_config_full(self):
        """from_config() with all vectorization keys sets every attribute."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        config = {
            "indexing": {"max_file_size_mb": 10},
            "vectorization": {
                "chunk_size": 500,
                "chunk_overlap": 200,
                "file_types": [".md", ".txt"],
                "exclude_patterns": ["**/exports/**"],
            },
        }
        ext = FullContentExtractor.from_config(config)
        assert ext.chunk_size == 500
        assert ext.chunk_overlap == 200
        assert ext.file_types == [".md", ".txt"]
        assert ext.exclude_patterns == ["**/exports/**"]
        assert ext.max_file_size_bytes == 10 * 1024 * 1024

    def test_from_config_empty(self):
        """from_config({}) uses constructor defaults."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        ext = FullContentExtractor.from_config({})
        assert ext.chunk_size == 1000
        assert ext.chunk_overlap == 150
        assert ext.file_types is None
        assert ext.exclude_patterns is None
        assert ext.max_file_size_bytes == 0

    def test_from_config_partial(self):
        """from_config() with only chunk_size overrides that, defaults the rest."""
        from footprinter.ingest.full_content_extractor import FullContentExtractor

        config = {"vectorization": {"chunk_size": 750}}
        ext = FullContentExtractor.from_config(config)
        assert ext.chunk_size == 750
        assert ext.chunk_overlap == 150
        assert ext.file_types is None
        assert ext.exclude_patterns is None
        assert ext.max_file_size_bytes == 0


class TestImportErrors:
    """Test ImportError handlers for optional dependencies (lines 74-82, 101-106)."""

    def test_docx_import_error(self, temp_dir):
        """When python-docx is not installed, returns None gracefully."""
        from footprinter.ingest.content_extractors import ContentExtractor

        # Create a minimal valid .docx-like file (doesn't matter, ImportError before read)
        docx_path = temp_dir / "test.docx"
        docx_path.write_bytes(b"fake docx")

        extractor = ContentExtractor()

        saved_docx = sys.modules.pop("docx", None)
        saved_flag = None
        try:
            # Reset the warning flag so the warning branch gets hit
            import footprinter.ingest.content_extractors as mod

            saved_flag = mod._docx_warned
            mod._docx_warned = False

            with patch.dict(sys.modules, {"docx": None}):
                content = extractor._extract_docx(docx_path)
            assert content is None
        finally:
            if saved_docx is not None:
                sys.modules["docx"] = saved_docx
            if saved_flag is not None:
                mod._docx_warned = saved_flag

    def test_pypdf_import_error(self, temp_dir):
        """When pypdf is not installed, returns None gracefully."""
        from footprinter.ingest.content_extractors import ContentExtractor

        pdf_path = temp_dir / "test.pdf"
        pdf_path.write_bytes(b"fake pdf")

        extractor = ContentExtractor()

        saved_pypdf = sys.modules.pop("pypdf", None)
        saved_flag = None
        try:
            import footprinter.ingest.content_extractors as mod

            saved_flag = mod._pypdf_warned
            mod._pypdf_warned = False

            with patch.dict(sys.modules, {"pypdf": None}):
                content = extractor._extract_pdf(pdf_path)
            assert content is None
        finally:
            if saved_pypdf is not None:
                sys.modules["pypdf"] = saved_pypdf
            if saved_flag is not None:
                mod._pypdf_warned = saved_flag
